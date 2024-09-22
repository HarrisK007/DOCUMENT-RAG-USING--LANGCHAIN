import fitz  
from docx import Document
import json
import pandas as pd
import os

# Define allowed file extensions
ALLOWED_EXTENSIONS = {'txt', 'csv', 'json', 'pdf', 'docx'}

# Document metadata variables
DOCUMENT_TYPE = None
DOCUMENT_SIZE = None
DOCUMENT_LINES = None
DOCUMENT_WORDS = None
DOCUMENT_PAGES = None

def allowed_file(filename):
    """Check if the file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def read_text_file(file_path):
    """Read and return content from a text file."""
    with open(file_path, 'r') as file:
        content = file.read()
    return content

def read_csv_file(file_path):
    """Read and return content from a CSV file."""
    df = pd.read_csv(file_path)
    return df.to_string()

def read_json_file(file_path):
    """Read and return content from a JSON file."""
    with open(file_path, 'r') as file:
        data = json.load(file)
    return json.dumps(data, indent=4)

def read_pdf_file(file_path):
    """Read and return content from a PDF file."""
    doc = fitz.open(file_path)
    text = ''
    global DOCUMENT_PAGES
    DOCUMENT_PAGES = doc.page_count  # Set number of pages for PDF
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text += page.get_text()
    return text

def read_docx_file(file_path):
    """Read and return content from a DOCX file."""
    doc = Document(file_path)
    text = [para.text for para in doc.paragraphs]
    
    # Approximate number of pages by number of paragraphs (can vary)
    global DOCUMENT_PAGES
    DOCUMENT_PAGES = len(doc.paragraphs)
    
    return '\n'.join(text)

def extract_text_from_file(file_path):
    """Extract text from a file and update document metadata."""
    global DOCUMENT_TYPE, DOCUMENT_SIZE, DOCUMENT_LINES, DOCUMENT_WORDS, DOCUMENT_PAGES
    
    ext = os.path.splitext(file_path)[1].lower()
    DOCUMENT_TYPE = ext.lstrip('.').upper()  # Set document type
    
    # Calculate file size
    DOCUMENT_SIZE = os.path.getsize(file_path)

    # Extract content based on the file type
    if ext == '.txt':
        content = read_text_file(file_path)
        DOCUMENT_PAGES = 1  # Assume 1 page for text files
    elif ext == '.csv':
        content = read_csv_file(file_path)
        DOCUMENT_PAGES = 1  # Assume 1 page for CSV files
    elif ext == '.json':
        content = read_json_file(file_path)
        DOCUMENT_PAGES = 1  # Assume 1 page for JSON files
    elif ext == '.pdf':
        content = read_pdf_file(file_path)
    elif ext == '.docx':
        content = read_docx_file(file_path)
    else:
        raise ValueError("Unsupported file format")
    
    # Set metadata for lines and words
    DOCUMENT_LINES = len(content.splitlines())
    DOCUMENT_WORDS = len(content.split())

    return content

def doc_info():
    """Return the current document's metadata."""
    return {
        "Document Type": DOCUMENT_TYPE,
        "Document Size (bytes)": DOCUMENT_SIZE//1000,
        "Document Lines": DOCUMENT_LINES,
        "Document Words": DOCUMENT_WORDS,
        "Document Pages": DOCUMENT_PAGES
    }
